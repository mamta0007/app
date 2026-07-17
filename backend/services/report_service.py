from fastapi import HTTPException
from models.analysis import Analysis
from models.resume import Resume
from models.jd import Jd
from utils.llm import llm
from models.roadmap import RoadMap
from utils.json_parser import parse_llm_json
from fastapi.responses import   FileResponse
from models.report import Report
from docx import Document
import os


def run_report(current_user, db):
    
    analysis = (
    db.query(Analysis)
    .filter(Analysis.user_id == current_user.id)
    .order_by(Analysis.id.desc())
    .first()
)
    if not analysis:
        raise HTTPException(404, "Analysis not found")
    
    
    resume = (
    db.query(Resume)
    .filter(Resume.id == analysis.resume_id)
    .first()
)
    if not resume:
        raise HTTPException(404, "Resume not found")
    jd = (
    db.query(Jd)
    .filter(Jd.id == analysis.jd_id)
    .first()
)
    if not jd:
        raise HTTPException(404, "JD not found")
    
    
    roadmap = (
    db.query(RoadMap)
    .filter(RoadMap.analysis_id == analysis.id)
    .first()
)

    if not roadmap:
        raise HTTPException(404, "Roadmap not found")
    
    existing_report = (
    db.query(Report)
    .filter(Report.analysis_id == analysis.id)
    .first()
)

    if existing_report:
        return FileResponse(
            path=existing_report.file_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=existing_report.file_name
        )
    
    template = """
Summarize and generate preparation tips.

Resume:
{resume_text}

JD:
{jd_text}

STRICT:
- Return ONLY valid JSON
- No explanation
- No markdown
Give JSON:
{{
  "resume_summary": "",
  "jd_summary": "",
  "tips": ""
}}
"""
    prompt = template.format(
    resume_text=resume.content,
        jd_text=jd.content
    )
    
    
    # LLM Response
    result = llm.invoke(prompt)
    
    try:
        response = parse_llm_json(result.content)
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="Failed to parse AI response."
        )
        
    # Final Report
    report_data = {
        "resume_summary": response.get("resume_summary", ""),
        "jd_summary": response.get("jd_summary", ""),
        "match_score": analysis.match_score,
        "skill_gap": analysis.missing_skills,
        "matching_skills":analysis.matching_skills,
        "learning_plan": roadmap.learning_plan,
        "tips": response.get("tips", "")
    }
    
    os.makedirs("reports", exist_ok=True)
    filename = f"report_{analysis.id}_{current_user.id}.docx"
    filepath = os.path.join("reports", filename)
    
    document=Document()
    
    document.add_heading("AI Interview Preparation Assistant (AIPA)", level=1)
    
    document.add_heading("Resume Summary", level=2)
    document.add_paragraph(report_data["resume_summary"])
    
    document.add_heading("Job Description Summary", level=2)
    document.add_paragraph(report_data["jd_summary"])

    document.add_heading("Match Score", level=2)
    document.add_paragraph(str(report_data["match_score"]))

    document.add_heading("Matching Skills", level=2)
    matching = report_data["matching_skills"]
    
    if isinstance(matching, list):
        for skill in matching:
            document.add_paragraph(skill, style="List Bullet")
    else:
        document.add_paragraph(str(matching))
    
    document.add_heading("Skill Gap", level=2)
    missing = report_data["skill_gap"]

    if isinstance(missing, list):
        for skill in missing:
            document.add_paragraph(skill, style="List Bullet")
    else:
        document.add_paragraph(str(missing))
        
    document.add_heading("Learning Plan", level=2)

    learning_plan = report_data["learning_plan"]

    if isinstance(learning_plan, dict):
        for week, task in learning_plan.items():
            document.add_paragraph(
                f"{week.replace('_', ' ').title()}: {task}",
                style="List Bullet"
            )
    else:
        document.add_paragraph(str(learning_plan))

    document.add_heading("Preparation Tips", level=2)
    document.add_paragraph(report_data["tips"])

    document.save(filepath)
                
    report = Report(
    user_id=current_user.id,
    analysis_id=analysis.id,
    file_name=filename,
    file_path=filepath
)
    db.add(report)
    db.commit()
    db.refresh(report)
    
            
    return FileResponse(
        path=filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )
    
   
        






